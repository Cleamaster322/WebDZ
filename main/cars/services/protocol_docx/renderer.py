import os
import re
import tempfile
from pathlib import Path

from PIL import Image, ImageOps
from docx import Document
from docx.shared import Inches


# =========================
# Regex patterns
# =========================

PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")
TR_IF_RE = re.compile(r"{%\s*tr\s+if\s+(.+?)\s*%}")
TR_ENDIF_RE = re.compile(r"{%\s*tr\s+endif\s*%}")


# =========================
# Image placeholders
# =========================

IMAGE_PLACEHOLDERS = {
    "photo_stand_test",
    "photo_gas_test",
    "photo_noise_test",
}


# =========================
# Image size limits for DOCX
# =========================

# Максимальные границы изображения в Word.
# Фото не обрезается и не вставляется в белый canvas.
# Оно пропорционально вписывается в эти границы.
#
# Регулируй только эти две величины.
MAX_IMAGE_WIDTH_INCHES = 3.2
MAX_IMAGE_HEIGHT_INCHES = 2.4


# =========================
# Paragraph / table iterators
# =========================

def iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        yield from iter_table_paragraphs(table)

    for section in document.sections:
        header = section.header
        footer = section.footer

        for paragraph in header.paragraphs:
            yield paragraph

        for table in header.tables:
            yield from iter_table_paragraphs(table)

        for paragraph in footer.paragraphs:
            yield paragraph

        for table in footer.tables:
            yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph

            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def iter_document_tables(document):
    for table in document.tables:
        yield table

    for section in document.sections:
        header = section.header
        footer = section.footer

        for table in header.tables:
            yield table

        for table in footer.tables:
            yield table


# =========================
# Conditional table rows
# =========================

def get_row_text(row):
    parts = []

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            parts.append(paragraph.text)

    return "\n".join(parts)


def delete_table_row(row):
    tr = row._tr
    tr.getparent().remove(tr)


def context_value_to_bool(value):
    """
    Приводит значение из context к boolean.

    Нормально работает для:
    True / False
    1 / 0
    "true" / "false"
    "yes" / "no"
    "1" / "0"
    """
    if value is True:
        return True

    if value is False:
        return False

    if value is None:
        return False

    if isinstance(value, (int, float)):
        return value != 0

    text = str(value).strip().lower()

    if text in ["", "0", "false", "none", "null", "no", "нет", "не применяется", "-"]:
        return False

    return True


def evaluate_tr_condition(expression, context):
    """
    Минимальный обработчик условий для таблиц.

    Поддерживает:
    {%tr if parking_light_present %}
    {%tr if not parking_light_present %}

    Специально не используем eval, чтобы не выполнять произвольный код.
    """
    expression = str(expression).strip()

    if expression.startswith("not "):
        variable_name = expression[4:].strip()
        return not context_value_to_bool(context.get(variable_name))

    return context_value_to_bool(context.get(expression))


def process_table_conditions(table, context):
    """
    Обрабатывает условные строки таблицы вида:

    {%tr if parking_light_present %}
        строка или несколько строк
    {%tr endif %}

    {%tr if not parking_light_present %}
        строка или несколько строк
    {%tr endif %}

    Служебные строки с {%tr ... %} всегда удаляются.
    Строки внутри блока удаляются, если условие False.
    """
    rows = list(table.rows)
    condition_stack = []
    rows_to_delete = []

    for row in rows:
        row_text = get_row_text(row)

        if_match = TR_IF_RE.search(row_text)
        endif_match = TR_ENDIF_RE.search(row_text)

        if if_match:
            parent_active = all(condition_stack) if condition_stack else True
            condition_result = evaluate_tr_condition(if_match.group(1), context)
            condition_stack.append(parent_active and condition_result)

            rows_to_delete.append(row)
            continue

        if endif_match:
            if condition_stack:
                condition_stack.pop()

            rows_to_delete.append(row)
            continue

        if condition_stack and not all(condition_stack):
            rows_to_delete.append(row)

    for row in rows_to_delete:
        try:
            delete_table_row(row)
        except Exception:
            pass

    # После удаления строк обрабатываем вложенные таблицы в оставшихся строках.
    for row in list(table.rows):
        for cell in row.cells:
            for nested_table in cell.tables:
                process_table_conditions(nested_table, context)


def process_document_conditions(document, context):
    for table in iter_document_tables(document):
        process_table_conditions(table, context)


# =========================
# Placeholder replacement
# =========================

def build_run_ranges(paragraph):
    ranges = []
    cursor = 0

    for run in paragraph.runs:
        start = cursor
        end = start + len(run.text)
        ranges.append((run, start, end))
        cursor = end

    return ranges


def find_run_index(ranges, char_index):
    for index, (_, start, end) in enumerate(ranges):
        if start <= char_index < end:
            return index

    if ranges and char_index == ranges[-1][2]:
        return len(ranges) - 1

    return None


def replace_placeholder_in_paragraph(paragraph, context):
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)

    if "{{" not in full_text:
        return

    matches = list(PLACEHOLDER_RE.finditer(full_text))

    if not matches:
        return

    # Идём с конца, чтобы позиции предыдущих matches не ломались.
    for match in reversed(matches):
        key = match.group(1)
        value = context.get(key, "")

        ranges = build_run_ranges(paragraph)

        start_index = find_run_index(ranges, match.start())
        end_index = find_run_index(ranges, match.end() - 1)

        if start_index is None or end_index is None:
            continue

        first_run, first_start, first_end = ranges[start_index]
        last_run, last_start, last_end = ranges[end_index]

        before = first_run.text[: match.start() - first_start]
        after = last_run.text[match.end() - last_start:]

        is_image = key in IMAGE_PLACEHOLDERS and value

        if start_index == end_index:
            if is_image:
                first_run.text = before + after
                try_add_picture(first_run, value)
            else:
                first_run.text = before + str(value) + after
        else:
            if is_image:
                first_run.text = before
                try_add_picture(first_run, value)
            else:
                first_run.text = before + str(value)

            for index in range(start_index + 1, end_index):
                paragraph.runs[index].text = ""

            last_run.text = after


# =========================
# Image processing
# =========================

def prepare_image_for_word(image_path):
    """
    Подготавливает изображение для вставки в Word:
    - исправляет поворот по EXIF;
    - убирает прозрачность;
    - сохраняет во временный JPEG.

    Важно:
    здесь нет обрезки, нет белого canvas и нет фиксированного прямоугольника.
    Фото сохраняется целиком.
    """
    source_path = Path(image_path)

    if not source_path.exists():
        return None

    try:
        image = Image.open(source_path)
        image = ImageOps.exif_transpose(image)

        if image.mode in ("RGBA", "LA"):
            background = Image.new("RGB", image.size, "white")
            background.paste(image, mask=image.split()[-1])
            image = background
        else:
            image = image.convert("RGB")

        temp_file = tempfile.NamedTemporaryFile(
            suffix=".jpg",
            delete=False,
        )

        temp_path = temp_file.name
        temp_file.close()

        image.save(temp_path, "JPEG", quality=92)

        return temp_path

    except Exception:
        return None


def calculate_docx_image_size(image_path):
    """
    Вычисляет размер изображения для Word.

    Логика:
    - исходные пропорции сохраняются;
    - фото не обрезается;
    - фото не искажается;
    - если фото горизонтальное, оно в первую очередь упирается в max width;
    - если фото вертикальное, оно в первую очередь упирается в max height;
    - итоговый размер не выходит за MAX_IMAGE_WIDTH_INCHES и MAX_IMAGE_HEIGHT_INCHES.

    Пример:
    1000x500  -> условно 3.2x1.6
    500x1000  -> условно 1.2x2.4
    """
    try:
        image = Image.open(image_path)
        image = ImageOps.exif_transpose(image)

        width_px, height_px = image.size

        if width_px <= 0 or height_px <= 0:
            return MAX_IMAGE_WIDTH_INCHES, MAX_IMAGE_HEIGHT_INCHES

        aspect_ratio = width_px / height_px

        if width_px >= height_px:
            # Горизонтальное или квадратное фото:
            # сначала растягиваем до максимальной ширины.
            width_inches = MAX_IMAGE_WIDTH_INCHES
            height_inches = width_inches / aspect_ratio

            # Если по высоте не помещается, ограничиваем по высоте.
            if height_inches > MAX_IMAGE_HEIGHT_INCHES:
                height_inches = MAX_IMAGE_HEIGHT_INCHES
                width_inches = height_inches * aspect_ratio
        else:
            # Вертикальное фото:
            # сначала растягиваем до максимальной высоты.
            height_inches = MAX_IMAGE_HEIGHT_INCHES
            width_inches = height_inches * aspect_ratio

            # Если по ширине не помещается, ограничиваем по ширине.
            if width_inches > MAX_IMAGE_WIDTH_INCHES:
                width_inches = MAX_IMAGE_WIDTH_INCHES
                height_inches = width_inches / aspect_ratio

        return width_inches, height_inches

    except Exception:
        return MAX_IMAGE_WIDTH_INCHES, MAX_IMAGE_HEIGHT_INCHES


def try_add_picture(run, image_path):
    prepared_path = prepare_image_for_word(image_path)

    if not prepared_path:
        return

    try:
        width_inches, height_inches = calculate_docx_image_size(prepared_path)

        run.add_picture(
            prepared_path,
            width=Inches(width_inches),
            height=Inches(height_inches),
        )
    except Exception:
        return
    finally:
        try:
            os.remove(prepared_path)
        except Exception:
            pass


# =========================
# Main render
# =========================

def render_protocol_docx(template_path, output_path, context):
    document = Document(template_path)

    # 1. Сначала обрабатываем условные строки таблиц:
    #    {%tr if ... %} / {%tr endif %}
    process_document_conditions(document, context)

    # 2. Потом заменяем обычные {{ placeholders }}
    for paragraph in iter_paragraphs(document):
        replace_placeholder_in_paragraph(paragraph, context)

    document.save(output_path)

    return output_path