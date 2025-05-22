from io import BytesIO
from docx import Document
from docx.shared import Pt

def create_car_word_doc(data: dict) -> BytesIO:
    """
    Создает Word-документ на основе данных и возвращает объект BytesIO.
    """

    doc = Document()

    # Заголовок
    doc.add_heading('Информация об автомобиле', level=1)

    # Основные данные
    doc.add_paragraph(f"Марка: {data.get('brand', '')}")
    doc.add_paragraph(f"Модель: {data.get('model', '')}")
    doc.add_paragraph(f"Поколение: {data.get('generation', '')}")
    doc.add_paragraph(f"Период поколения: {data.get('generation_period', '')}")
    doc.add_paragraph(f"Комплектация: {data.get('configuration', '')}")
    doc.add_paragraph(f"Название двигателя: {data.get('engine_name', '')}")

    doc.add_heading('Технические данные', level=2)

    car_data = data.get('car_data', {})
    if car_data:
        for key, value in car_data.items():
            # Преобразуем ключ к более читаемому виду, если нужно
            # Например "front_tires" -> "Передние шины"
            readable_key = key.replace('_', ' ').capitalize()
            doc.add_paragraph(f"{readable_key}: {value}")
    else:
        doc.add_paragraph("Данные отсутствуют.")

    # Сохраняем документ в память
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
